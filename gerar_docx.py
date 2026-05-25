from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Margens ──────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(3)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2)

# ── Helpers ──────────────────────────────────────────
def add_heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_paragraph(text, bold=False, italic=False, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def page_break():
    doc.add_page_break()

AZUL = (30, 58, 138)

# ════════════════════════════════════════════════════
# CAPA
# ════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('CLUBE DE DESBRAVADORES THIAGO WHITE')
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(*AZUL)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Documentação do Sistema de Gerenciamento')
run.font.size = Pt(15)
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('─' * 45)
run.font.color.rgb = RGBColor(*AZUL)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

for linha in [
    ('Associação:', 'Igreja Adventista do Sétimo Dia'),
    ('Local:', 'Fortaleza — CE'),
    ('Tecnologia:', 'Next.js · TypeScript · Supabase'),
    ('Versão:', '1.0.0'),
    ('Data:', datetime.date.today().strftime('%d/%m/%Y')),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f'{linha[0]}  ')
    r1.bold = True
    r1.font.size = Pt(12)
    r2 = p.add_run(linha[1])
    r2.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Fortaleza — 2026')
run.font.size = Pt(11)
run.italic = True

page_break()

# ════════════════════════════════════════════════════
# SUMÁRIO (manual)
# ════════════════════════════════════════════════════
add_heading('Sumário', 1, AZUL)

sumario = [
    ('1', 'Introdução', '3'),
    ('2', 'Justificativa', '4'),
    ('3', 'Objetivos', '5'),
    ('3.1', 'Objetivo Geral', '5'),
    ('3.2', 'Objetivos Específicos', '5'),
    ('4', 'Descrição do Sistema', '6'),
    ('4.1', 'Visão Geral', '6'),
    ('4.2', 'Stack Tecnológica', '6'),
    ('4.3', 'Módulos do Sistema', '7'),
    ('5', 'Módulo: Dashboard', '8'),
    ('6', 'Módulo: Cadastros', '9'),
    ('7', 'Módulo: Mensalidades', '11'),
    ('8', 'Módulo: Caixa', '12'),
    ('9', 'Módulo: Custos', '13'),
    ('10', 'Módulo: Patrimônio', '14'),
    ('11', 'Módulo: Atas', '15'),
    ('12', 'Módulo: Atos', '16'),
    ('13', 'Módulo: Relatórios', '17'),
    ('14', 'Banco de Dados', '18'),
    ('15', 'Testes do Sistema', '20'),
    ('16', 'Página Estática (index.html)', '21'),
    ('17', 'Conclusão', '22'),
]

table = doc.add_table(rows=len(sumario), cols=2)
table.style = 'Table Grid'
for i, (num, titulo, pag) in enumerate(sumario):
    row = table.rows[i]
    row.cells[0].text = f'{num}  {titulo}'
    row.cells[1].text = pag
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)
        shade_cell(cell, 'F8FAFC' if i % 2 == 0 else 'FFFFFF')
    row.cells[0].width = Cm(13)
    row.cells[1].width = Cm(2)

page_break()

# ════════════════════════════════════════════════════
# 1. INTRODUÇÃO
# ════════════════════════════════════════════════════
add_heading('1. Introdução', 1, AZUL)

add_paragraph(
    'O presente documento descreve o Sistema de Gerenciamento do Clube de Desbravadores '
    'Thiago White, desenvolvido para centralizar e automatizar as atividades administrativas '
    'do clube. O sistema foi concebido para atender às necessidades da diretoria, secretaria '
    'e tesouraria, substituindo planilhas e controles manuais por uma plataforma web '
    'moderna, acessível de qualquer dispositivo com navegador.'
)
add_paragraph(
    'Os Desbravadores são uma organização da Igreja Adventista do Sétimo Dia voltada para '
    'crianças e adolescentes entre 10 e 15 anos, com foco no desenvolvimento espiritual, '
    'educacional e social. A gestão eficiente do clube é essencial para garantir o '
    'cumprimento da missão e o bem-estar de todos os membros.'
)
add_paragraph(
    'Este documento está organizado em seções que cobrem desde a justificativa e os '
    'objetivos do projeto até a descrição técnica de cada módulo, o modelo de banco de '
    'dados e os resultados dos testes realizados.'
)

page_break()

# ════════════════════════════════════════════════════
# 2. JUSTIFICATIVA
# ════════════════════════════════════════════════════
add_heading('2. Justificativa', 1, AZUL)

add_paragraph(
    'A gestão de um Clube de Desbravadores envolve diversas responsabilidades simultâneas: '
    'controle de membros, cobrança de mensalidades, registro de reuniões, gestão financeira, '
    'inventário de patrimônio e produção de documentos oficiais. Na ausência de um sistema '
    'dedicado, essas tarefas costumam ser realizadas de forma fragmentada — em planilhas '
    'isoladas, cadernos físicos e documentos dispersos — gerando retrabalho, perda de '
    'informação e dificuldade na tomada de decisão.'
)
add_paragraph(
    'A criação de um sistema integrado justifica-se pelos seguintes problemas identificados '
    'na rotina do clube:'
)

bullets_justificativa = [
    'Controle manual de mensalidades sujeito a erros e esquecimentos.',
    'Ausência de histórico consolidado das finanças do clube.',
    'Dificuldade em localizar documentos (atas, atos, autorizações) rapidamente.',
    'Inventário de patrimônio sem atualização periódica.',
    'Falta de visão gerencial para tomada de decisões pela diretoria.',
    'Geração manual e demorada de relatórios para prestação de contas.',
]
for b in bullets_justificativa:
    add_bullet(b)

doc.add_paragraph()
add_paragraph(
    'Diante desse cenário, a implantação de um sistema web centralizado representa uma '
    'solução prática, acessível e de baixo custo operacional, capaz de transformar a '
    'gestão do clube e liberar tempo da equipe voluntária para o que realmente importa: '
    'o trabalho com os jovens.'
)

page_break()

# ════════════════════════════════════════════════════
# 3. OBJETIVOS
# ════════════════════════════════════════════════════
add_heading('3. Objetivos', 1, AZUL)

add_heading('3.1 Objetivo Geral', 2)
add_paragraph(
    'Desenvolver um sistema web de gerenciamento completo para o Clube de Desbravadores '
    'Thiago White, integrando em uma única plataforma o controle de membros, finanças, '
    'patrimônio e documentação administrativa.'
)

add_heading('3.2 Objetivos Específicos', 2)
bullets_obj = [
    'Cadastrar e gerenciar os desbravadores e seus dados pessoais.',
    'Controlar o pagamento mensal de cada membro com visualização em grade.',
    'Registrar todas as entradas e saídas financeiras do caixa do clube.',
    'Gerenciar despesas operacionais (custos) separadamente do caixa.',
    'Manter um inventário atualizado dos bens pertencentes ao clube.',
    'Registrar formalmente as reuniões em atas e os documentos em atos administrativos.',
    'Gerar relatórios para impressão e prestação de contas.',
    'Oferecer um dashboard com indicadores gerenciais em tempo real.',
    'Garantir segurança dos dados com Row-Level Security (RLS) no banco de dados.',
    'Cobrir o sistema com testes automatizados para garantir qualidade do software.',
]
for b in bullets_obj:
    add_bullet(b)

page_break()

# ════════════════════════════════════════════════════
# 4. DESCRIÇÃO DO SISTEMA
# ════════════════════════════════════════════════════
add_heading('4. Descrição do Sistema', 1, AZUL)

add_heading('4.1 Visão Geral', 2)
add_paragraph(
    'O sistema é uma aplicação web de página única (SPA) desenvolvida com Next.js 15 '
    'e App Router. Todas as telas são renderizadas no lado do cliente ("use client"), '
    'comunicando-se diretamente com o banco de dados Supabase via sua API REST. '
    'A interface utiliza Tailwind CSS para estilização e componentes Radix UI para '
    'elementos interativos acessíveis.'
)
add_paragraph(
    'O acesso ao sistema é feito por navegador web, sem necessidade de instalação. '
    'Em ambiente de produção, o banco de dados fica hospedado no Supabase Cloud; '
    'em desenvolvimento local, utiliza-se Docker com o Supabase CLI.'
)

add_heading('4.2 Stack Tecnológica', 2)

headers = ['Camada', 'Tecnologia', 'Versão', 'Função']
rows_stack = [
    ['Front-end',    'Next.js + React',    '15 / 19',  'Framework web com App Router'],
    ['Linguagem',    'TypeScript',         '5.x',      'Tipagem estática'],
    ['Estilo',       'Tailwind CSS',       '4.x',      'Utilitários CSS'],
    ['Componentes',  'Radix UI',           '2.x',      'UI acessível e sem estilo próprio'],
    ['Ícones',       'Lucide React',       '—',        'Biblioteca de ícones SVG'],
    ['Gráficos',     'Recharts',           '2.x',      'Gráficos de barra e pizza'],
    ['Banco de dados','Supabase (PostgreSQL)','—',     'BaaS com RLS e API REST'],
    ['Autenticação', 'Supabase Auth',      '—',        'Controle de acesso por roles'],
    ['Testes',       'Vitest + RTL',       '3.x / 16.x','Testes unitários e de componente'],
    ['Package manager','pnpm',             '9.x',      'Gerenciador de pacotes'],
]

t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
hdr = t.rows[0].cells
for i, h in enumerate(headers):
    hdr[i].text = h
    shade_cell(hdr[i], '1E3A8A')
    for para in hdr[i].paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)

for j, row_data in enumerate(rows_stack):
    row = t.add_row().cells
    for i, val in enumerate(row_data):
        row[i].text = val
        shade_cell(row[i], 'EFF6FF' if j % 2 == 0 else 'FFFFFF')
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

add_heading('4.3 Módulos do Sistema', 2)
add_paragraph('O sistema é composto por 9 módulos principais:')

modulos = [
    ('Dashboard',    'Visão geral com indicadores e gráficos gerenciais.'),
    ('Cadastros',    'Desbravadores, Unidades, Classes e Especialidades.'),
    ('Mensalidades', 'Grade de controle de pagamentos mensais por membro.'),
    ('Caixa',        'Registro de entradas e saídas financeiras.'),
    ('Custos',       'Controle de despesas operacionais do clube.'),
    ('Patrimônio',   'Inventário de bens com estado de conservação.'),
    ('Atas',         'Registro formal das reuniões e assembleias.'),
    ('Atos',         'Documentos administrativos (resoluções, portarias, etc.).'),
    ('Relatórios',   'Geração de seis tipos de relatórios para impressão.'),
]
for nome, desc in modulos:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'{nome}: ')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(desc)
    r2.font.size = Pt(11)

page_break()

# ════════════════════════════════════════════════════
# 5. DASHBOARD
# ════════════════════════════════════════════════════
add_heading('5. Módulo: Dashboard', 1, AZUL)

add_paragraph(
    'O Dashboard é a tela inicial do sistema e oferece uma visão executiva do clube em '
    'tempo real. Todos os dados são carregados automaticamente ao acessar a página.'
)

add_heading('Indicadores (cards)', 2)
headers2 = ['Indicador', 'Descrição']
rows2 = [
    ['Desbravadores Ativos',       'Total de membros com status "ativo".'],
    ['Inadimplentes (mês atual)',   'Ativos sem mensalidade paga no mês corrente.'],
    ['Saldo de Caixa',             'Total de entradas menos total de saídas.'],
    ['Patrimônio Total',           'Soma dos valores de aquisição (exceto baixados).'],
]
t2 = doc.add_table(rows=1, cols=2)
t2.style = 'Table Grid'
for i, h in enumerate(headers2):
    t2.rows[0].cells[i].text = h
    shade_cell(t2.rows[0].cells[i], '1E3A8A')
    for para in t2.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255,255,255)
for j, rd in enumerate(rows2):
    row = t2.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        shade_cell(row[i], 'EFF6FF' if j%2==0 else 'FFFFFF')
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()
add_heading('Gráficos', 2)
add_bullet('Mensalidades por Mês: gráfico de barras (pagos × pendentes) para cada mês do ano.')
add_bullet('Distribuição por Unidade: gráfico de pizza com a proporção de membros por unidade.')
add_bullet('Últimas Transações: tabela com as movimentações de caixa mais recentes.')

page_break()

# ════════════════════════════════════════════════════
# 6. CADASTROS
# ════════════════════════════════════════════════════
add_heading('6. Módulo: Cadastros', 1, AZUL)

add_paragraph(
    'O módulo de Cadastros é o coração do sistema e abrange quatro sub-módulos: '
    'Desbravadores, Unidades, Classes e Especialidades.'
)

add_heading('6.1 Desbravadores', 2)
add_paragraph(
    'Gerencia todos os membros do clube. Cada desbravador possui uma ficha completa '
    'com dados pessoais, vínculo com unidade e classe, e histórico de especialidades.'
)
add_paragraph('Campos do cadastro:', bold=True)

headers3 = ['Campo', 'Tipo', 'Obrigatório']
rows3 = [
    ['Nome completo',       'Texto',          'Sim'],
    ['CPF',                 'Texto',          'Não'],
    ['RG',                  'Texto',          'Não'],
    ['Data de nascimento',  'Data',           'Não'],
    ['Sexo',                'masculino / feminino', 'Não'],
    ['Cargo',               'desbravador / conselheiro / instrutor / aspirante / diretor', 'Sim'],
    ['Status',              'ativo / inativo / transferido / desligado', 'Sim'],
    ['Unidade',             'Vínculo com tabela unidades', 'Não'],
    ['Classe atual',        'Vínculo com tabela classes', 'Não'],
    ['Telefone / Email',    'Texto',          'Não'],
    ['Endereço',            'Texto',          'Não'],
    ['Nome do responsável', 'Texto',          'Não'],
    ['Data de ingresso',    'Data',           'Não'],
]
t3 = doc.add_table(rows=1, cols=3)
t3.style = 'Table Grid'
for i, h in enumerate(headers3):
    t3.rows[0].cells[i].text = h
    shade_cell(t3.rows[0].cells[i], '1E3A8A')
    for para in t3.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255,255,255)
for j, rd in enumerate(rows3):
    row = t3.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        shade_cell(row[i], 'EFF6FF' if j%2==0 else 'FFFFFF')
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

doc.add_paragraph()
add_heading('6.2 Unidades', 2)
add_paragraph('Grupos internos do clube (ex.: Águias, Falcões, Leões). Campos: nome, líder e contato do líder.')

add_heading('6.3 Classes', 2)
add_paragraph('Níveis de progressão pré-cadastrados pelo sistema.')

headers4 = ['Classe', 'Nível', 'Faixa etária']
rows4 = [
    ['Amigo','1','10 anos'], ['Companheiro','2','11 anos'],
    ['Pesquisador','3','12 anos'], ['Pioneiro','4','13 anos'],
    ['Excursionista','5','14 anos'], ['Guia','6','15 anos'],
]
t4 = doc.add_table(rows=1, cols=3)
t4.style = 'Table Grid'
for i, h in enumerate(headers4):
    t4.rows[0].cells[i].text = h
    shade_cell(t4.rows[0].cells[i], '1E3A8A')
    for para in t4.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255,255,255)
for j, rd in enumerate(rows4):
    row = t4.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        shade_cell(row[i], 'EFF6FF' if j%2==0 else 'FFFFFF')
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()
add_heading('6.4 Especialidades', 2)
add_paragraph('Cadastro das especialidades disponíveis para conquista. Campos: nome, área, nível (1–3) e descrição.')

page_break()

# ════════════════════════════════════════════════════
# 7. MENSALIDADES
# ════════════════════════════════════════════════════
add_heading('7. Módulo: Mensalidades', 1, AZUL)

add_paragraph(
    'O módulo de mensalidades exibe uma grade onde cada linha é um desbravador ativo '
    'e cada coluna é um mês do ano. O operador registra ou cancela pagamentos com um '
    'único clique na célula correspondente.'
)

add_heading('Legenda da grade', 2)
headers5 = ['Símbolo', 'Cor', 'Significado']
rows5 = [
    ['✓',  'Verde',    'Mensalidade paga'],
    ['—',  'Cinza',    'Pendente (não paga)'],
    ['!',  'Vermelho', 'Atrasada (prazo vencido)'],
    ['I',  'Azul',     'Isento'],
]
t5 = doc.add_table(rows=1, cols=3)
t5.style = 'Table Grid'
for i, h in enumerate(headers5):
    t5.rows[0].cells[i].text = h
    shade_cell(t5.rows[0].cells[i], '1E3A8A')
    for para in t5.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255,255,255)
for j, rd in enumerate(rows5):
    row = t5.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        shade_cell(row[i], 'EFF6FF' if j%2==0 else 'FFFFFF')
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(
    'O cabeçalho de cada coluna exibe o placar do mês (ex.: "4/8" = 4 pagos de 8 ativos). '
    'É possível filtrar por unidade e navegar entre anos com as setas de ano.'
)
add_paragraph(
    'O valor padrão de cada mensalidade é R$ 30,00 e a data de pagamento é registrada '
    'automaticamente no momento do clique.'
)

page_break()

# ════════════════════════════════════════════════════
# 8. CAIXA
# ════════════════════════════════════════════════════
add_heading('8. Módulo: Caixa', 1, AZUL)

add_paragraph(
    'Registra todas as movimentações financeiras do clube. Os três cards no topo '
    'exibem o total de Entradas, o total de Saídas e o Saldo atual.'
)

add_heading('Campos de um lançamento', 2)
headers6 = ['Campo', 'Tipo', 'Obrigatório']
rows6 = [
    ['Tipo',       'entrada / saída',                            'Sim'],
    ['Descrição',  'Texto livre',                                'Sim'],
    ['Valor',      'Decimal (R$)',                               'Sim'],
    ['Data',       'Data',                                       'Sim'],
    ['Categoria',  'mensalidade / evento / doação / material / outro', 'Sim'],
    ['Referência', 'Texto (nome do evento ou pessoa)',           'Não'],
]
t6 = doc.add_table(rows=1, cols=3)
t6.style = 'Table Grid'
for i, h in enumerate(headers6):
    t6.rows[0].cells[i].text = h
    shade_cell(t6.rows[0].cells[i], '1E3A8A')
    for para in t6.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255,255,255)
for j, rd in enumerate(rows6):
    row = t6.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        shade_cell(row[i], 'EFF6FF' if j%2==0 else 'FFFFFF')
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(
    'É possível filtrar os lançamentos por tipo (entrada/saída). '
    'A exclusão de lançamentos requer confirmação e é permanente.'
)

page_break()

# ════════════════════════════════════════════════════
# 9. CUSTOS
# ════════════════════════════════════════════════════
add_heading('9. Módulo: Custos', 1, AZUL)

add_paragraph(
    'Registra as despesas operacionais do clube separadamente do caixa, '
    'permitindo melhor organização interna e acompanhamento do status de pagamento.'
)

headers7 = ['Campo', 'Tipo', 'Obrigatório']
rows7 = [
    ['Descrição',          'Texto',                                   'Sim'],
    ['Valor',              'Decimal (R$)',                            'Sim'],
    ['Data',               'Data de vencimento ou realização',        'Sim'],
    ['Categoria',          'material / evento / manutenção / alimentação / transporte / outro', 'Sim'],
    ['Fornecedor',         'Texto',                                   'Não'],
    ['Nota Fiscal',        'Número da NF',                            'Não'],
    ['Status pagamento',   'pendente / pago',                        'Sim'],
]
t7 = doc.add_table(rows=1, cols=3)
t7.style = 'Table Grid'
for i, h in enumerate(headers7):
    t7.rows[0].cells[i].text = h
    shade_cell(t7.rows[0].cells[i], '1E3A8A')
    for para in t7.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255,255,255)
for j, rd in enumerate(rows7):
    row = t7.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        shade_cell(row[i], 'EFF6FF' if j%2==0 else 'FFFFFF')
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(
    'O resumo do topo exibe os totais de despesas Pendentes e Pagas, '
    'facilitando a visualização do compromisso financeiro do clube.'
)

page_break()

# ════════════════════════════════════════════════════
# 10. PATRIMÔNIO
# ════════════════════════════════════════════════════
add_heading('10. Módulo: Patrimônio', 1, AZUL)

add_paragraph(
    'Mantém o inventário de todos os bens pertencentes ao clube: equipamentos, '
    'mobiliário, instrumentos, materiais e outros itens de valor.'
)

headers8 = ['Campo', 'Tipo', 'Obrigatório']
rows8 = [
    ['Nome do bem',          'Texto',                  'Sim'],
    ['Nº de tombamento',     'Texto único',            'Não'],
    ['Descrição',            'Texto',                  'Não'],
    ['Valor de aquisição',   'Decimal (R$)',           'Não'],
    ['Data de aquisição',    'Data',                   'Não'],
    ['Estado de conservação','ótimo/bom/regular/ruim/baixado', 'Sim'],
    ['Localização',          'Texto',                  'Não'],
    ['Observações',          'Texto',                  'Não'],
]
t8 = doc.add_table(rows=1, cols=3)
t8.style = 'Table Grid'
for i, h in enumerate(headers8):
    t8.rows[0].cells[i].text = h
    shade_cell(t8.rows[0].cells[i], '1E3A8A')
    for para in t8.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255,255,255)
for j, rd in enumerate(rows8):
    row = t8.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        shade_cell(row[i], 'EFF6FF' if j%2==0 else 'FFFFFF')
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()
add_paragraph(
    'Bens com estado "Baixado" são desconsiderados no cálculo do Patrimônio Total '
    'exibido no Dashboard.'
)

page_break()

# ════════════════════════════════════════════════════
# 11. ATAS
# ════════════════════════════════════════════════════
add_heading('11. Módulo: Atas', 1, AZUL)

add_paragraph(
    'Registra formalmente as reuniões e assembleias do clube, garantindo rastreabilidade '
    'das decisões tomadas. As atas são numeradas sequencialmente e podem ser marcadas '
    'como aprovadas após deliberação na reunião seguinte.'
)

headers9 = ['Campo', 'Tipo', 'Obrigatório']
rows9 = [
    ['Número',          'Inteiro sequencial',                          'Sim'],
    ['Data',            'Data da reunião',                             'Sim'],
    ['Tipo',            'diretoria / clube / pais / extraordinária',   'Sim'],
    ['Local',           'Texto',                                       'Sim'],
    ['Pauta',           'Assuntos tratados',                           'Sim'],
    ['Conteúdo',        'Texto completo da ata',                      'Sim'],
    ['Aprovada',        'Booleano (Sim/Não)',                          'Sim'],
    ['Data de aprovação','Data',                                       'Não'],
]
t9 = doc.add_table(rows=1, cols=3)
t9.style = 'Table Grid'
for i, h in enumerate(headers9):
    t9.rows[0].cells[i].text = h
    shade_cell(t9.rows[0].cells[i], '1E3A8A')
    for para in t9.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255,255,255)
for j, rd in enumerate(rows9):
    row = t9.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        shade_cell(row[i], 'EFF6FF' if j%2==0 else 'FFFFFF')
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

page_break()

# ════════════════════════════════════════════════════
# 12. ATOS
# ════════════════════════════════════════════════════
add_heading('12. Módulo: Atos', 1, AZUL)

add_paragraph(
    'Registra os documentos administrativos formais emitidos pela direção do clube: '
    'resoluções, portarias, circulares e ofícios. Cada ato é numerado '
    'sequencialmente e identificado por tipo e assinante.'
)

headers10 = ['Campo', 'Tipo', 'Obrigatório']
rows10 = [
    ['Número',    'Inteiro sequencial',                             'Sim'],
    ['Data',      'Data de emissão',                               'Sim'],
    ['Tipo',      'resolução / portaria / circular / ofício',       'Sim'],
    ['Assunto',   'Título do documento',                           'Sim'],
    ['Conteúdo',  'Texto completo',                                'Sim'],
    ['Assinante', 'Nome do responsável pela assinatura',           'Sim'],
]
t10 = doc.add_table(rows=1, cols=3)
t10.style = 'Table Grid'
for i, h in enumerate(headers10):
    t10.rows[0].cells[i].text = h
    shade_cell(t10.rows[0].cells[i], '1E3A8A')
    for para in t10.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255,255,255)
for j, rd in enumerate(rows10):
    row = t10.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        shade_cell(row[i], 'EFF6FF' if j%2==0 else 'FFFFFF')
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

page_break()

# ════════════════════════════════════════════════════
# 13. RELATÓRIOS
# ════════════════════════════════════════════════════
add_heading('13. Módulo: Relatórios', 1, AZUL)

add_paragraph(
    'O módulo de relatórios centraliza a geração de documentos para impressão ou '
    'exportação em PDF. Cada relatório é acessado por um card na tela e pode ser '
    'impresso com Ctrl+P ou salvo como PDF pelo navegador.'
)

headers11 = ['Relatório', 'Descrição', 'Público-alvo']
rows11 = [
    ['Autorização de Saída', 'Formulário de autorização para responsáveis assinarem antes de eventos externos.', 'Responsáveis / Direção'],
    ['Fluxo de Caixa',       'Extrato financeiro com entradas, saídas e saldo por período.', 'Tesouraria'],
    ['Patrimônio',           'Lista completa do inventário de bens com valores e estados.', 'Direção'],
    ['Livro Ata e Atos',     'Compilado de atas e atos numerados em formato de livro.', 'Secretaria'],
    ['Mensalidades',         'Boletim de pagamentos por mês e por desbravador.', 'Tesouraria'],
    ['Cadastros',            'Listagem de desbravadores com dados de contato, unidade e classe.', 'Secretaria'],
]
t11 = doc.add_table(rows=1, cols=3)
t11.style = 'Table Grid'
for i, h in enumerate(headers11):
    t11.rows[0].cells[i].text = h
    shade_cell(t11.rows[0].cells[i], '1E3A8A')
    for para in t11.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255,255,255)
for j, rd in enumerate(rows11):
    row = t11.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        shade_cell(row[i], 'EFF6FF' if j%2==0 else 'FFFFFF')
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

page_break()

# ════════════════════════════════════════════════════
# 14. BANCO DE DADOS
# ════════════════════════════════════════════════════
add_heading('14. Banco de Dados', 1, AZUL)

add_paragraph(
    'O banco de dados é PostgreSQL gerenciado pelo Supabase. Todas as tabelas possuem '
    'Row-Level Security (RLS) ativada, garantindo que somente usuários autorizados '
    'acessem os dados. Duas migrations compõem o schema.'
)

add_heading('14.1 Tabelas', 2)

headers12 = ['Tabela', 'Descrição', 'Campos principais']
rows12 = [
    ['unidades',                  'Grupos internos do clube',         'id, nome, lider_nome, ativo'],
    ['classes',                   'Níveis de progressão',             'id, nome, nivel (1–6)'],
    ['especialidades',            'Especialidades disponíveis',       'id, nome, area, nivel (1–3)'],
    ['desbravadores',             'Membros do clube',                 'id, nome, cpf, cargo, status, unidade_id, classe_id'],
    ['desbravador_especialidades','Vínculo membro × especialidade',   'desbravador_id, especialidade_id, status, data_conclusao'],
    ['mensalidades',              'Pagamentos mensais',               'desbravador_id, mes_referencia, valor, status, data_pagamento'],
    ['caixa_transacoes',          'Movimentações financeiras',        'id, tipo, descricao, valor, data, categoria'],
    ['custos',                    'Despesas operacionais',            'id, descricao, valor, data, categoria, status_pagamento'],
    ['patrimonio',                'Inventário de bens',               'id, nome, numero_tombamento, valor_aquisicao, estado'],
    ['atas',                      'Registros de reuniões',            'id, numero, data, tipo, local, conteudo, aprovada'],
    ['atos',                      'Documentos administrativos',       'id, numero, data, tipo, assunto, conteudo, assinante'],
]
t12 = doc.add_table(rows=1, cols=3)
t12.style = 'Table Grid'
for i, h in enumerate(headers12):
    t12.rows[0].cells[i].text = h
    shade_cell(t12.rows[0].cells[i], '1E3A8A')
    for para in t12.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255,255,255)
for j, rd in enumerate(rows12):
    row = t12.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        shade_cell(row[i], 'EFF6FF' if j%2==0 else 'FFFFFF')
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(9)

doc.add_paragraph()
add_heading('14.2 Migrations', 2)

headers13 = ['Arquivo', 'Descrição']
rows13 = [
    ['0001_clube_thiago_white.sql', 'Schema completo: tabelas, enums, RLS para role authenticated e classes pré-cadastradas.'],
    ['0002_anon_access.sql',        'Políticas RLS para role anon — usadas em desenvolvimento local sem autenticação.'],
]
t13 = doc.add_table(rows=1, cols=2)
t13.style = 'Table Grid'
for i, h in enumerate(headers13):
    t13.rows[0].cells[i].text = h
    shade_cell(t13.rows[0].cells[i], '1E3A8A')
    for para in t13.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255,255,255)
for j, rd in enumerate(rows13):
    row = t13.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        shade_cell(row[i], 'EFF6FF' if j%2==0 else 'FFFFFF')
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

page_break()

# ════════════════════════════════════════════════════
# 15. TESTES
# ════════════════════════════════════════════════════
add_heading('15. Testes do Sistema', 1, AZUL)

add_paragraph(
    'O sistema conta com uma suíte de testes automatizados usando Vitest e '
    'React Testing Library (RTL), com ambiente jsdom que simula o navegador. '
    'Os mocks do Supabase são construídos com vi.hoisted() para garantir que '
    'os dados fictícios estejam disponíveis antes do carregamento dos módulos.'
)

add_heading('Resultado da suíte de testes', 2)

headers14 = ['Arquivo de teste', 'Módulo testado', 'Casos de teste']
rows14 = [
    ['dashboard.test.tsx',      'Dashboard',    '4 — indicadores, gráficos, transações'],
    ['desbravadores.test.tsx',  'Desbravadores','4 — listagem, busca, botão novo'],
    ['mensalidades.test.tsx',   'Mensalidades', '4 — grade, células, placar'],
    ['caixa.test.tsx',          'Caixa',        '4 — saldo, lançamentos, filtro'],
    ['custos.test.tsx',         'Custos',       '4 — listagem, resumo pendente/pago'],
    ['patrimonio.test.tsx',     'Patrimônio',   '4 — listagem, estado conservação'],
    ['atas.test.tsx',           'Atas',         '4 — listagem, aprovada, data'],
    ['atos.test.tsx',           'Atos',         '3 — listagem, tipo, assunto'],
    ['utils.test.ts',           'Utilitários',  '8 — formatação, cálculos'],
    ['types.test.ts',           'Tipos/Enums',  '6 — labels e constantes'],
]
t14 = doc.add_table(rows=1, cols=3)
t14.style = 'Table Grid'
for i, h in enumerate(headers14):
    t14.rows[0].cells[i].text = h
    shade_cell(t14.rows[0].cells[i], '1E3A8A')
    for para in t14.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255,255,255)
for j, rd in enumerate(rows14):
    row = t14.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        shade_cell(row[i], 'EFF6FF' if j%2==0 else 'FFFFFF')
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

doc.add_paragraph()

p = doc.add_paragraph()
r = p.add_run('Resultado final: 51 testes passando em 10 arquivos — 100% de aprovação.')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(22, 163, 74)

page_break()

# ════════════════════════════════════════════════════
# 16. PÁGINA ESTÁTICA
# ════════════════════════════════════════════════════
add_heading('16. Página Estática (index.html)', 1, AZUL)

add_paragraph(
    'Além do sistema web, foi desenvolvida uma página estática em HTML puro '
    '(index.html) na raiz do repositório. O arquivo pode ser aberto diretamente '
    'no navegador sem servidor, instalação ou conexão com banco de dados.'
)

add_heading('Elementos da página', 2)

headers15 = ['Elemento', 'Descrição']
rows15 = [
    ['Título (title)',       'Aparece na aba do navegador.'],
    ['Cabeçalho (header)',   'Faixa azul com nome e localização do clube.'],
    ['Menu esquerdo (nav)',  'Links para seções da página e módulos do sistema.'],
    ['Corpo (main)',         'Banner 1, Banner 2, cards de indicadores, Sobre e Valores.'],
    ['Banner 1',             'Fundo que pulsa alternando de azul para roxo continuamente (CSS @keyframes).'],
    ['Banner 2',             'Texto com avisos deslizando da direita para a esquerda (marquee CSS).'],
    ['Rodapé (footer)',      'Faixa escura com copyright e ano dinâmico via JavaScript.'],
]
t15 = doc.add_table(rows=1, cols=2)
t15.style = 'Table Grid'
for i, h in enumerate(headers15):
    t15.rows[0].cells[i].text = h
    shade_cell(t15.rows[0].cells[i], '1E3A8A')
    for para in t15.rows[0].cells[i].paragraphs:
        for run in para.runs:
            run.bold = True; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255,255,255)
for j, rd in enumerate(rows15):
    row = t15.add_row().cells
    for i, v in enumerate(rd):
        row[i].text = v
        shade_cell(row[i], 'EFF6FF' if j%2==0 else 'FFFFFF')
        for para in row[i].paragraphs:
            for run in para.runs:
                run.font.size = Pt(10)

page_break()

# ════════════════════════════════════════════════════
# 17. CONCLUSÃO
# ════════════════════════════════════════════════════
add_heading('17. Conclusão', 1, AZUL)

add_paragraph(
    'O Sistema de Gerenciamento do Clube de Desbravadores Thiago White foi desenvolvido '
    'com sucesso, atendendo a todos os objetivos propostos. A plataforma centraliza em '
    'um único ambiente web as principais atividades administrativas do clube, que '
    'anteriormente dependiam de controles manuais fragmentados.'
)
add_paragraph(
    'A solução entregue oferece uma interface limpa e intuitiva, acessível de qualquer '
    'dispositivo com navegador. O Dashboard proporciona à diretoria uma visão gerencial '
    'imediata, enquanto os módulos de Mensalidades e Caixa reduzem significativamente '
    'o risco de erros e perdas de informação na área financeira.'
)
add_paragraph(
    'Do ponto de vista técnico, a escolha do Supabase como backend eliminou a necessidade '
    'de desenvolver uma API própria, acelerando o desenvolvimento e garantindo segurança '
    'nativa por meio de Row-Level Security. O uso de TypeScript em toda a base de código '
    'reduziu bugs de tipo em tempo de desenvolvimento, enquanto a suíte de 51 testes '
    'automatizados garante a estabilidade do sistema em futuras evoluções.'
)
add_paragraph(
    'O projeto demonstra que é possível construir um sistema de gestão completo, seguro '
    'e testado utilizando tecnologias modernas e gratuitas, tornando a solução acessível '
    'para clubes de qualquer porte, independentemente de orçamento de TI.'
)

add_heading('Possibilidades de evolução futura', 2)
bullets_evolucao = [
    'Implementação de autenticação por e-mail para acesso seguro à diretoria.',
    'Notificações automáticas por WhatsApp ou e-mail para mensalidades em atraso.',
    'Aplicativo mobile com React Native para acesso offline em acampamentos.',
    'Módulo de eventos com controle de presença e lista de inscritos.',
    'Integração com sistemas financeiros para conciliação bancária automática.',
    'Exportação de relatórios em formato PDF com assinatura digital.',
]
for b in bullets_evolucao:
    add_bullet(b)

doc.add_paragraph()
add_paragraph(
    'O sistema está disponível publicamente no repositório GitHub '
    '(github.com/italoarruda/clube-thiago-white) sob licença MIT, '
    'podendo ser utilizado e adaptado por outros clubes de desbravadores.',
    italic=True
)

# ── Salvar ────────────────────────────────────────
caminho = '/Users/italoarruda/Documents/GitHub/clube-thiago-white/docs/Documentacao_Sistema_Clube_Thiago_White.docx'
doc.save(caminho)
print(f'Documento salvo em: {caminho}')
